import streamlit as st
import pandas as pd
import io
import re
import zipfile
import spacy
from docx import Document
import fitz  # PyMuPDF
import firebase_admin
from firebase_admin import credentials, firestore


# --- 1. ІНІЦІАЛІЗАЦІЯ FIREBASE ---
def init_firebase():
    if not firebase_admin._apps:
        try:
            fb_secrets = dict(st.secrets["firebase"])
            fb_secrets["private_key"] = fb_secrets["private_key"].replace("\\n", "\n")
            cred = credentials.Certificate(fb_secrets)
            firebase_admin.initialize_app(cred)
        except Exception as e:
            st.error(f"Firebase Error: {e}")
            st.stop()
    return firestore.client()


db = init_firebase()


# --- 2. ЗАВАНТАЖЕННЯ МОДЕЛІ NLP ---
@st.cache_resource
def load_nlp():
    return spacy.load("pl_core_news_sm")


nlp = load_nlp()

# --- 3. ЛОГІКА ТОКЕНІЗАЦІЇ ---

WHITELIST = {
    "Biuro", "Miejsce", "Data", "Dyrektor", "Prezes", "Firma", "Konto",
    "Numer", "Adres", "Spotkanie", "Protokół", "Raport", "Warszawa", "Poznań"
}


def get_token(value, entity_type, mapping):
    if value not in mapping:
        idx = len([v for v in mapping.values() if entity_type in v]) + 1
        mapping[value] = f"<{entity_type}_{idx}>"
    return mapping[value]


def process_text(text, custom_list, mapping):
    if not text or str(text) == "nan":
        return text

    doc = nlp(str(text))
    result_text = str(text)
    matches = []

    # 1. Покращений Regex (PESEL, NIP, Sp. z o.o., Решта)
    patterns = {
        "EMAIL": r"[\w\.-]+@[\w\.-]+\.\w+",
        "IBAN": r"\b[A-Z]{2}\d{2}[A-Z0-9]{12,30}\b",
        "CARD": r"\b(?:\d[ -]*?){13,16}\b",
        "PHONE": r"(?:\+\d{2,3}|0)\s?[\d\-\s]{7,12}\b",
        "PESEL": r"\b\d{11}\b",
        "NIP": r"\b\d{10}\b|\b\d{3}-\d{3}-\d{2}-\d{2}\b",
        # Спеціальний пошук компаній з юридичними формами (Sp. z o.o., S.A. і т.д.)
        "ORG_LEGAL": r"\b[A-Z][\w\s\.-]+(?:Sp\.\s?z\s?o\.o\.|S\.A\.|Sp\.\s?k\.|Sp\.\s?j\.)\b"
    }

    for label, pattern in patterns.items():
        for match in re.finditer(pattern, result_text):
            # Якщо це юридична форма, маркуємо як ORG
            final_label = "ORG" if label == "ORG_LEGAL" else label
            matches.append((match.start(), match.end(), final_label))

    # 2. Custom Blacklist з БД
    if custom_list:
        for word in custom_list:
            if len(word) > 2:
                for m in re.finditer(re.escape(word), result_text, re.IGNORECASE):
                    matches.append((m.start(), m.end(), "CUSTOM"))

    # 3. NLP (тільки вільні зони)
    for ent in doc.ents:
        if ent.text in WHITELIST:
            continue

        is_overlap = any(m[0] < ent.end_char and ent.start_char < m[1] for m in matches)

        if not is_overlap and ent.label_ in ["PERSON", "ORG", "GPE"]:
            if len(ent.text) > 2 and ent.text[0].isupper():
                matches.append((ent.start_char, ent.end_char, ent.label_))

    # Сортування з кінця
    matches = sorted(list(set(matches)), key=lambda x: x[0], reverse=True)

    for start, end, label in matches:
        original = result_text[start:end]
        token = get_token(original, label, mapping)
        result_text = result_text[:start] + token + result_text[end:]

    return result_text


# --- 4. ОБРОБНИКИ ФАЙЛІВ ---

def handle_pdf(file_bytes, custom_list, mapping):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page in doc:
        # Для PDF додаємо пошук Sp. z o.o.
        patterns_to_redact = [r"\S+@\S+", r"[A-Z]{2}\d{2}[A-Z0-9]{12,30}", r"Sp\.\s?z\s?o\.o\."]
        for pattern in patterns_to_redact:
            for m in page.search_for(pattern):
                page.add_redact_annot(m, fill=(0, 0, 0))

        if custom_list:
            for word in custom_list:
                for area in page.search_for(word):
                    token = get_token(word, "CUSTOM", mapping)
                    page.add_redact_annot(area, fill=(0, 0, 0))
                    page.apply_redactions()
                    page.insert_text(area.tl, token, color=(1, 1, 1), fontsize=8)
    return doc.tobytes()


def handle_docx(file_bytes, custom_list, mapping):
    doc = Document(io.BytesIO(file_bytes))
    for p in doc.paragraphs:
        p.text = process_text(p.text, custom_list, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = process_text(cell.text, custom_list, mapping)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def handle_xlsx(file_bytes, custom_list, mapping):
    df = pd.read_excel(io.BytesIO(file_bytes))
    for col in df.columns:
        df[col] = df[col].astype(str).apply(lambda x: process_text(x, custom_list, mapping))
    out = io.BytesIO()
    df.to_excel(out, index=False)
    return out.getvalue()


# --- 5. ІНТЕРФЕЙС ---

st.set_page_config(page_title="Secure Tokenizer Pro", page_icon="🛡️")
st.title("🛡️ Secure Tokenizer Pro")

with st.sidebar:
    st.header("Ustawienia")
    ref = db.collection("settings").document("blacklist")
    db_data = ref.get().to_dict() if ref.get().exists else {"names": ""}
    db_names = db_data.get("names", "")

    input_names = st.text_area("Czarna lista (nazwy firm, nazwiska):", value=db_names)
    if st.button("Zapisz w Firebase"):
        ref.set({"names": input_names})
        st.success("Zapisano!")

    blacklist = [n.strip() for n in input_names.split(",") if n.strip()]

files = st.file_uploader("Dodaj pliki (PDF, DOCX, XLSX)", accept_multiple_files=True, type=['pdf', 'docx', 'xlsx'])

if files and st.button("Uruchom anonimizację"):
    session_mapping = {}
    zip_buf = io.BytesIO()

    with zipfile.ZipFile(zip_buf, 'w') as zf:
        for f in files:
            data = f.read()
            ext = f.name.split(".")[-1].lower()

            if ext == 'pdf':
                res = handle_pdf(data, blacklist, session_mapping)
            elif ext == 'docx':
                res = handle_docx(data, blacklist, session_mapping)
            elif ext == 'xlsx':
                res = handle_xlsx(data, blacklist, session_mapping)

            zf.writestr(f"anonymized_{f.name}", res)

        report = pd.DataFrame(session_mapping.items(), columns=["Oryginał", "Token"])
        rep_buf = io.BytesIO()
        report.to_excel(rep_buf, index=False)
        zf.writestr("raport_mapowania.xlsx", rep_buf.getvalue())

    st.success("Przetwarzanie zakończone!")
    st.download_button("Pobierz wyniki (ZIP)", zip_buf.getvalue(), "wyniki.zip")